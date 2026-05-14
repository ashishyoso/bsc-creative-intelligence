"""Briefs API (Module 5).

Endpoints:
- POST /briefs/from-formula    → generate a brief from a Magic Formula intent
- GET  /briefs                 → list briefs
- GET  /briefs/{id}            → detail
- PATCH /briefs/{id}           → edit (title, status, brief_markdown, etc.)
- DELETE /briefs/{id}          → delete
- POST /briefs/{id}/ship       → mark shipped (lifecycle helper)
"""
from __future__ import annotations

import json
from datetime import datetime

from fastapi import APIRouter, Body, Depends, HTTPException
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from app.db.models import AuditEvent, Brief, BriefStatus
from app.db.session import get_db
from app.patterns.brief_generator import (
    assemble_brief_markdown,
    generate_brief_content,
)
from app.patterns.formula import generate_formula

router = APIRouter(prefix="/briefs", tags=["briefs"])


class FromFormulaRequest(BaseModel):
    target_sku: str
    metric: str = "roas"
    persona: str | None = None
    pain_addressed: str | None = None
    audio_language: str | None = None
    format_constraint: str | None = None
    awareness_stage: str | None = None
    forbidden_archetypes: list[str] = Field(default_factory=list)
    min_n: int = 3
    title_override: str | None = None


@router.post("/from-formula")
def from_formula(req: FromFormulaRequest, db: Session = Depends(get_db)):
    """Generate a brief from a Magic Formula intent. One call = one new draft brief."""
    formula = generate_formula(
        db,
        target_sku=req.target_sku,
        metric=req.metric,
        persona=req.persona,
        audio_language=req.audio_language,
        format_constraint=req.format_constraint,
        forbidden_archetypes=req.forbidden_archetypes,
        min_n=req.min_n,
    )
    if formula.cohort_size == 0:
        raise HTTPException(400, "no_cohort: no historical creatives match these constraints")

    generated = generate_brief_content(formula)
    if "error" in generated:
        raise HTTPException(500, f"hook_generation_failed: {generated['error']}")

    title = req.title_override or generated.get("title") or f"{req.target_sku} — auto-generated brief"
    markdown = assemble_brief_markdown(formula, generated)

    brief = Brief(
        title=title,
        status=BriefStatus.DRAFT.value,
        target_sku=req.target_sku,
        target_metric=req.metric,
        persona=req.persona,
        pain_addressed=req.pain_addressed,
        awareness_stage=req.awareness_stage,
        audio_language=req.audio_language,
        format_constraint=req.format_constraint,
        formula_json=json.dumps({
            "recommendations": [r.__dict__ for r in formula.recommendations],
            "references": [r.__dict__ for r in formula.references],
            "risks": formula.risks,
        }, default=str),
        overall_confidence=formula.overall_confidence,
        cohort_size=formula.cohort_size,
        brief_markdown=markdown,
        verbal_hooks_json=json.dumps(generated.get("verbal_hooks", [])),
        on_screen_hooks_json=json.dumps(generated.get("on_screen_hooks", [])),
        mechanic=generated.get("mechanic"),
        music_direction=generated.get("music_direction"),
        talent_direction=generated.get("talent_direction"),
        duration_target_seconds=generated.get("duration_target_seconds"),
        reference_asset_ids=",".join(r.asset_id for r in formula.references),
        forbidden_archetypes_json=json.dumps(req.forbidden_archetypes),
    )
    db.add(brief)
    db.flush()
    db.add(AuditEvent(
        action_type="brief_created",
        target_entity="brief",
        target_id=str(brief.id),
        payload=f"sku={req.target_sku} cost=Rs.{generated.get('_cost_inr', 0)}",
    ))
    db.commit()

    return _serialize_brief(brief)


def _serialize_brief(brief: Brief) -> dict:
    return {
        "id": brief.id,
        "title": brief.title,
        "status": brief.status,
        "owner": brief.owner,
        "target_sku": brief.target_sku,
        "target_metric": brief.target_metric,
        "persona": brief.persona,
        "pain_addressed": brief.pain_addressed,
        "awareness_stage": brief.awareness_stage,
        "audio_language": brief.audio_language,
        "format_constraint": brief.format_constraint,
        "overall_confidence": brief.overall_confidence,
        "cohort_size": brief.cohort_size,
        "brief_markdown": brief.brief_markdown,
        "verbal_hooks": json.loads(brief.verbal_hooks_json) if brief.verbal_hooks_json else [],
        "on_screen_hooks": json.loads(brief.on_screen_hooks_json) if brief.on_screen_hooks_json else [],
        "mechanic": brief.mechanic,
        "music_direction": brief.music_direction,
        "talent_direction": brief.talent_direction,
        "duration_target_seconds": brief.duration_target_seconds,
        "reference_asset_ids": brief.reference_asset_ids.split(",") if brief.reference_asset_ids else [],
        "formula_json": json.loads(brief.formula_json) if brief.formula_json else None,
        "notes": brief.notes,
        "created_at": brief.created_at.isoformat() if brief.created_at else None,
        "updated_at": brief.updated_at.isoformat() if brief.updated_at else None,
        "shipped_at": brief.shipped_at.isoformat() if brief.shipped_at else None,
    }


@router.get("")
def list_briefs(db: Session = Depends(get_db), status: str | None = None):
    q = db.query(Brief)
    if status:
        q = q.filter(Brief.status == status)
    out = []
    for b in q.order_by(Brief.created_at.desc()).all():
        out.append({
            "id": b.id,
            "title": b.title,
            "status": b.status,
            "target_sku": b.target_sku,
            "persona": b.persona,
            "overall_confidence": b.overall_confidence,
            "cohort_size": b.cohort_size,
            "created_at": b.created_at.isoformat() if b.created_at else None,
        })
    return out


@router.get("/{brief_id}")
def get_brief(brief_id: int, db: Session = Depends(get_db)):
    brief = db.get(Brief, brief_id)
    if not brief:
        raise HTTPException(404, "brief_not_found")
    return _serialize_brief(brief)


class BriefPatch(BaseModel):
    title: str | None = None
    status: str | None = None
    brief_markdown: str | None = None
    notes: str | None = None
    verbal_hooks: list[str] | None = None
    on_screen_hooks: list[str] | None = None
    mechanic: str | None = None
    music_direction: str | None = None
    talent_direction: str | None = None
    duration_target_seconds: int | None = None


@router.patch("/{brief_id}")
def edit_brief(brief_id: int, body: BriefPatch, db: Session = Depends(get_db)):
    brief = db.get(Brief, brief_id)
    if not brief:
        raise HTTPException(404, "brief_not_found")
    if body.title is not None:
        brief.title = body.title
    if body.status is not None:
        brief.status = body.status
        if body.status == BriefStatus.SHIPPED.value and not brief.shipped_at:
            brief.shipped_at = datetime.utcnow()
    if body.brief_markdown is not None:
        brief.brief_markdown = body.brief_markdown
    if body.notes is not None:
        brief.notes = body.notes
    if body.verbal_hooks is not None:
        brief.verbal_hooks_json = json.dumps(body.verbal_hooks)
    if body.on_screen_hooks is not None:
        brief.on_screen_hooks_json = json.dumps(body.on_screen_hooks)
    if body.mechanic is not None:
        brief.mechanic = body.mechanic
    if body.music_direction is not None:
        brief.music_direction = body.music_direction
    if body.talent_direction is not None:
        brief.talent_direction = body.talent_direction
    if body.duration_target_seconds is not None:
        brief.duration_target_seconds = body.duration_target_seconds
    db.add(AuditEvent(
        action_type="brief_edited",
        target_entity="brief",
        target_id=str(brief_id),
        payload=str(body.model_dump(exclude_unset=True)),
    ))
    db.commit()
    return _serialize_brief(brief)


@router.delete("/{brief_id}")
def delete_brief(brief_id: int, db: Session = Depends(get_db)):
    brief = db.get(Brief, brief_id)
    if not brief:
        raise HTTPException(404, "brief_not_found")
    db.delete(brief)
    db.commit()
    return {"ok": True}


@router.get("/{brief_id}/export.docx")
def export_brief_docx(brief_id: int, db: Session = Depends(get_db)):
    """US-5.5: export brief as a docx file."""
    from io import BytesIO
    from fastapi.responses import Response

    brief = db.get(Brief, brief_id)
    if not brief:
        raise HTTPException(404, "brief_not_found")

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise HTTPException(500, "python-docx not installed; pip install python-docx")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading(brief.title, level=0)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    def _kv(k, v):
        row = table.add_row().cells
        row[0].text = k
        row[1].text = str(v) if v is not None else "—"
    _kv("SKU", brief.target_sku)
    _kv("Optimizing for", brief.target_metric)
    _kv("Target persona", brief.persona)
    _kv("Language", brief.audio_language)
    _kv("Format", brief.format_constraint)
    _kv("Pain addressed", brief.pain_addressed)
    _kv("Awareness", brief.awareness_stage)
    _kv("Overall confidence", brief.overall_confidence)
    _kv("Cohort", f"{brief.cohort_size} historical creatives")

    doc.add_heading("Verbal hook options", level=1)
    verbal = json.loads(brief.verbal_hooks_json) if brief.verbal_hooks_json else []
    for i, hook in enumerate(verbal, 1):
        doc.add_paragraph(f"{i}. {hook}", style="List Number")

    doc.add_heading("On-screen text options", level=1)
    on_screen = json.loads(brief.on_screen_hooks_json) if brief.on_screen_hooks_json else []
    for i, hook in enumerate(on_screen, 1):
        doc.add_paragraph(f"{i}. {hook}", style="List Number")

    doc.add_heading("Production direction", level=1)
    _t = doc.add_table(rows=0, cols=2)
    _t.style = "Light Grid Accent 1"
    def _kv2(k, v):
        row = _t.add_row().cells
        row[0].text = k
        row[1].text = str(v) if v else "—"
    _kv2("Mechanic", brief.mechanic)
    _kv2("Music", brief.music_direction)
    _kv2("Talent", brief.talent_direction)
    _kv2("Duration target", f"{brief.duration_target_seconds}s" if brief.duration_target_seconds else None)

    doc.add_heading("Mandatory close", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Khul ke Khel Bro")
    run.bold = True

    if brief.formula_json:
        formula = json.loads(brief.formula_json)
        recs = formula.get("recommendations", [])
        doc.add_heading("Tag manifest (data-grounded)", level=1)
        tag_tbl = doc.add_table(rows=1, cols=4)
        tag_tbl.style = "Light Grid Accent 1"
        hdr = tag_tbl.rows[0].cells
        hdr[0].text = "Dimension"; hdr[1].text = "Value"; hdr[2].text = "N"; hdr[3].text = "Confidence"
        for r in recs:
            if not r.get("value"): continue
            row = tag_tbl.add_row().cells
            row[0].text = r.get("label", "")
            row[1].text = str(r.get("value", ""))
            row[2].text = str(r.get("n", ""))
            row[3].text = str(r.get("confidence", ""))

        risks = formula.get("risks", [])
        if risks:
            doc.add_heading("Risks", level=1)
            for r in risks:
                doc.add_paragraph(f"• {r}")

    if brief.notes:
        doc.add_heading("Internal notes", level=1)
        doc.add_paragraph(brief.notes)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_title = "".join(c for c in brief.title if c.isalnum() or c in (" ", "-", "_"))[:60].strip().replace(" ", "_")
    filename = f"{safe_title or 'brief'}_brief_{brief_id}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
